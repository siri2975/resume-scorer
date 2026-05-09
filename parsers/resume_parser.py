"""
resume_parser.py
----------------
Extracts raw text from uploaded resume files (PDF or DOCX).
Supports multi-page PDFs and complex DOCX formatting.
"""

import os
import pdfplumber
import docx


def extract_text_from_pdf(filepath: str) -> str:
    """Extract all text from a PDF file using pdfplumber."""
    text = ""
    try:
        with pdfplumber.open(filepath) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        raise ValueError(f"Failed to parse PDF: {str(e)}")
    return text.strip()


def extract_text_from_docx(filepath: str) -> str:
    """Extract all text from a DOCX file using python-docx."""
    text = ""
    try:
        doc = docx.Document(filepath)
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text.strip() + "\n"

        # Also extract text from tables inside DOCX
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text.strip() + " "
                text += "\n"
    except Exception as e:
        raise ValueError(f"Failed to parse DOCX: {str(e)}")
    return text.strip()


def parse_resume(filepath: str) -> str:
    """
    Main entry point.
    Auto-detects file type from extension and extracts text.

    Args:
        filepath: Absolute or relative path to the uploaded resume file.

    Returns:
        Raw extracted text as a single string.

    Raises:
        ValueError: If file type is unsupported or parsing fails.
    """
    if not os.path.exists(filepath):
        raise FileNotFoundError(f"File not found: {filepath}")

    ext = os.path.splitext(filepath)[1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(filepath)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(filepath)
    else:
        raise ValueError(f"Unsupported file type: '{ext}'. Only PDF and DOCX are supported.")